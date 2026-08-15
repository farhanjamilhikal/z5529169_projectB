"""Build and format the student-only authorship decision workbook."""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "student_review" / "AUTHORSHIP_AND_INVESTMENT_DECISION_GUIDE.md"
OUTPUT = ROOT / "student_review" / "AUTHORSHIP_AND_INVESTMENT_DECISION_WORKBOOK.docx"
OXBLOOD = RGBColor(0x6F, 0x1D, 0x2C)
BLUE = RGBColor(0x3B, 0x6F, 0x8F)


def _set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def format_workbook() -> None:
    document = Document(OUTPUT)
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        header = section.header.paragraphs[0]
        header.text = "Signal & Story | Student Review Workbook"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = BLUE

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.color.rgb = OXBLOOD
    by_id = {style.style_id: style for style in styles}
    for style_id, size, colour in (
        ("Heading1", 16, OXBLOOD),
        ("Heading2", 12, BLUE),
        ("Heading3", 10.5, BLUE),
    ):
        style = by_id[style_id]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = colour

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    for table in document.tables:
        table.autofit = True
        if table.rows:
            _set_repeat_header(table.rows[0])
        for row in table.rows:
            _prevent_row_split(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8.5)

    document.save(OUTPUT)


def main() -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise SystemExit("Pandoc is required to rebuild this optional student workbook.")
    subprocess.run(
        [pandoc, str(SOURCE), "--toc", "-o", str(OUTPUT)],
        cwd=ROOT,
        check=True,
    )
    format_workbook()
    print(OUTPUT)


if __name__ == "__main__":
    main()
