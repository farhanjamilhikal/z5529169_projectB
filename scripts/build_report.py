"""Build the editable prefinalised Word report from the verified Markdown draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "report" / "PREFINALISED_REPORT_DRAFT.md"
OUTPUT = ROOT / "report" / "report.docx"

INK = "242424"
OXBLOOD = "6F1D2C"
BLUE = "547A92"
IVORY = "F7F3EA"
PALE = "F1E8DD"
GREY = "8A8A86"
BRICK = "B54A42"

FIGURE_CAPTIONS = {
    "growth_of_one_comparison.png": "Figure 1. Out-of-sample growth of one dollar after 10 bps transaction costs. Crypto-only funds use the seven-day calendar; equity and combined funds use the equity decision calendar. Source: course-provided prices and student calculations.",
    "fund_drawdowns.png": "Figure 2. Live out-of-sample drawdowns for minimum-variance funds. Source: course-provided prices and student calculations.",
    "combined_weights_over_time.png": "Figure 3. Monthly combined-fund target weights, showing the eight largest average holdings and Other. Source: course-provided prices and student calculations.",
    "sector_sentiment_index.png": "Figure 4. Finance-augmented VADER sector sentiment, equal-weighted across tickers with no-news ticker-days neutral. Source: course-provided headlines and student calculations.",
    "fusion_before_after.png": "Figure 5. Growth and drawdown before and after reliability-gated sentiment fusion. Source: course-provided prices, headlines and student calculations.",
    "risk_return_across_funds.png": "Figure 6. Annualised out-of-sample return versus annualised volatility across funds and methods after transaction costs. Point size increases with the Sharpe ratio. Source: course-provided prices and student calculations.",
    "crypto_sleeve_floor_sensitivity.png": "Figure 7. Combined minimum-variance sensitivity to a 0, 10, 20 or 30 per cent minimum total cryptocurrency sleeve. The constrained variants are research tests, not offered funds. Source: course-provided prices and student calculations.",
}


def set_font(run, name="Arial", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, before, after in [
        ("Heading 1", 16, 18, 9),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(OXBLOOD if name != "Heading 3" else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.16


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=False):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_inline(paragraph, text, *, size=10.5, color=INK):
    text = (
        text.replace("—", "-")
        .replace("–", "-")
        .replace("‑", "-")
        .replace("*t*", "t")
        .replace("*i*", "i")
        .replace("*c*", "c")
        .replace("*z*", "z")
        .replace("*q*", "q")
    )
    pieces = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
    for piece in pieces:
        if not piece:
            continue
        bold = piece.startswith("**") and piece.endswith("**")
        code = piece.startswith("`") and piece.endswith("`")
        italic = piece.startswith("*") and piece.endswith("*") and not bold
        clean = piece[2:-2] if bold else piece[1:-1] if (code or italic) else piece
        run = paragraph.add_run(clean)
        set_font(
            run,
            "Consolas" if code else "Arial",
            size=9 if code else size,
            color=color,
            bold=bold,
            italic=italic,
        )


def add_callout(doc, text, review=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF2CC" if review else PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size=9.5, color=BRICK if review else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_table(doc, lines):
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0], *rows[2:]]
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    if ncols == 2:
        widths = [2800, 6560]
    elif ncols == 5:
        widths = [3600, 1440, 1440, 1200, 1680]
    elif ncols == 6:
        widths = [3000, 1250, 1250, 1000, 1400, 1460]
    else:
        widths = [9360 // ncols] * ncols
        widths[-1] += 9360 - sum(widths)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, OXBLOOD)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_inline(
                p, value, size=8.1 if ncols >= 5 else 9, color="FFFFFF" if r_idx == 0 else INK
            )
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Signal & Story | FINS3645 Project B | z5529169 | Page ")
    set_font(run, size=8, color=GREY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    # Editorial cover pattern with Signal & Story named palette override.
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(kicker, "FINS3645 PROJECT B", size=11, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("SIGNAL & STORY")
    set_font(r, "Georgia", 30, OXBLOOD, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(
        sub, "Systematic Multi-Asset Funds with News-Sentiment Analytics", size=15, color=BLUE
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(
        meta,
        "DFF Stations 3 and 4\nStudent ID z5529169\nSample 1 January 2020 to 31 December 2023",
        size=10,
        color=GREY,
    )
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(status, "SYSTEMATIC FUNDS | SENTIMENT | INVESTOR DASHBOARD", size=9.5, color=BRICK)
    doc.add_page_break()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Executive summary")
    lines = lines[start:]
    i = 0
    figures = []
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("!["):
            match = re.search(r"\((.*?)\)", line)
            if match:
                path = (ROOT / "report" / match.group(1)).resolve()
                figures.append(path)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(
                    p,
                    f"See {FIGURE_CAPTIONS[path.name].split('.')[0]} in Appendix C.",
                    size=9,
                    color=BLUE,
                )
            i += 1
            continue
        if line.startswith("**Figure "):
            i += 1
            continue
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and set(lines[i + 1].replace("|", "").replace(":", "").replace("-", "").strip())
            == set()
        ):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("> "):
            add_callout(doc, line[2:], review="Student authorship gate" in line)
        elif "[STUDENT REVIEW:" in line:
            add_callout(doc, line.replace("`", ""), review=True)
        elif line.startswith("- ["):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line == r"\[":
            equation = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                equation.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            add_inline(p, " ".join(equation), size=10, color=BLUE)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    doc.add_page_break()
    doc.add_heading("Appendix C. Required figures", level=1)
    for idx, path in enumerate(figures):
        if idx:
            doc.add_page_break()
        doc.add_picture(str(path), width=Inches(6.45))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cap, FIGURE_CAPTIONS[path.name], size=8.5, color=GREY)

    props = doc.core_properties
    props.title = "Signal & Story - FINS3645 Project B"
    props.subject = "FINS3645 Project B report"
    props.author = "z5529169"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
