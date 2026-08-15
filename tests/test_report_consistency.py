from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parent.parent


def _document_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        blocks.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_report_metrics_match_generated_csv():
    metrics = pd.read_csv(ROOT / "results/tables/performance_metrics.csv")
    text = _document_text(ROOT / "report/report.docx")
    for row in metrics.itertuples(index=False):
        assert row.fund in text
        for value in (
            row.annualised_return,
            row.annualised_volatility,
            row.maximum_drawdown,
        ):
            assert f"{value:.2%}" in text
        assert f"{row.sharpe_ratio:.2f}" in text


def test_report_contains_every_required_exhibit_without_review_placeholders():
    path = ROOT / "report/report.docx"
    document = Document(path)
    text = _document_text(path)
    assert len(document.inline_shapes) == 7
    for caption in (
        "Figure 1.",
        "Figure 2.",
        "Figure 3.",
        "Figure 4.",
        "Figure 5.",
        "Figure 6.",
    ):
        assert caption in text
    assert any(term in text.lower() for term in ("risk and return", "risk-return"))
    assert "[STUDENT REVIEW" not in text
    assert (ROOT / "report/report.pdf").stat().st_size > 500_000
